"""Extract text from word files stored in docx format."""

import io

from docx import Document


def extract_docx_content(file_content: bytes) -> str:
    """Extract text from a word file.

    Parameters
    ----------
    file_content : bytes
        The content of the word file.


    Returns
    -------
    str
        The extracted text.

    """
    with io.BytesIO(file_content) as file_like:
        try:
            doc = Document(file_like)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            return "\n\n".join(full_text)
        except Exception as e:
            error_msg = f"Failed to extract text from Word document: {e}"
            raise ValueError(error_msg) from e
